#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Generate a one‑page customised résumé, but now using Ollama
instead of OpenAI.

All functions that used the OpenAI API are replaced with
``call_ollama()``.  The prompts themselves are unchanged –
you can keep them identical or tweak them for your own model.
"""

import argparse
from pathlib import Path
import textwrap

# ------------------------------------------------------------------
# 1. Imports – Ollama client
# ------------------------------------------------------------------
try:
    from ollama import Client
except Exception as exc:
    raise RuntimeError(
        "The Ollama Python client is required. "
        "Install it with `pip install ollama`."
    ) from exc

OLLAMA_HOST = "http://localhost:11434"
ollama_client = Client(host=OLLAMA_HOST)

# ------------------------------------------------------------------
# 2. Helper for cleaning / loading documents
# ------------------------------------------------------------------
def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")

def read_pdf(path: Path) -> str:
    """
    Very light PDF‑to‑text fallback.  We keep the old
    implementation to keep the script portable.
    """
    from pdfminer.high_level import extract_text
    return extract_text(str(path))

def read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def clean_text(txt: str) -> str:
    return " ".join(txt.split())

# ------------------------------------------------------------------
# 3. The Ollama wrapper
# ------------------------------------------------------------------
def call_ollama(model: str, prompt: str, max_tokens: int = 512) -> str:
    """
    A tiny wrapper around Ollama’s ``client.generate`` call.
    ``stream=False`` returns the whole answer in a single dict.
    The returned value contains a key called ``'response'``.
    """
    response = ollama_client.generate(
        model=model,
        prompt=prompt,
        stream=False,
    )
    return response["response"].strip()

# ------------------------------------------------------------------
# 4. Pipeline components (same as before, but use Ollama)
# ------------------------------------------------------------------
def summarise_job_ad(job_ad_text: str, model: str) -> dict:
        prompt = textwrap.dedent(f"""
        You are a senior HR analyst.  
        Summarise this job posting in the following JSON format:

        {{
            "title": "...",
            "responsibilities": [...],
            "key_skills": [...]
        }}

        **Job posting**  
        {job_ad_text}
        """)
        raw = call_ollama(model, prompt)
        print("DEBUG: Model response for job summary:\n", raw)
        # Remove Markdown formatting if present
        lines = raw.splitlines()
        cleaned_lines = []
        for line in lines:
            if line.strip().startswith('```') or line.strip().lower() == 'json':
                continue
            cleaned_lines.append(line)
        cleaned_raw = '\n'.join(cleaned_lines).strip()
        import json
        return json.loads(cleaned_raw)

def extract_resume_skills(resume_text: str, model: str) -> list[str]:
    prompt = textwrap.dedent(f"""
    List all distinct technical skills mentioned in the résumé below,
    one per line.  If a skill appears multiple times, keep it only once.

    {resume_text}
    """)
    raw = call_ollama(model, prompt)
    # Ollama returns raw text – split into lines, strip empty ones
    skills = [line.strip() for line in raw.splitlines() if line.strip()]
    return sorted(set(skills), key=lambda x: x.lower())

def align_skills(job_skills: list[str], resume_skills: list[str]) -> list[tuple[str, bool]]:
    return [(skill, skill in resume_skills) for skill in job_skills]

def generate_custom_resume(job_summary: dict, original_resume: str,
                           alignment: list[tuple[str, bool]], model: str) -> str:
    skill_markers = "\n".join(
        f"- {skill} {'✔' if present else '❌'}"
        for skill, present in alignment
    )
    prompt = textwrap.dedent(f"""
    Keep the applicant's original résumé structure but emphasise the skills that match the job.
    Place them in a Key Skills section and reference the job title and core responsibilities in the objective.

    **Original résumé**  
    {original_resume}

    **Key skills**  
    {skill_markers}

    **Job title**: {job_summary["title"]}
    **Responsibilities**:
    {chr(10).join(f"- {b}" for b in job_summary["responsibilities"])}

    Produce a single‑page résumé (≈ 500 words).  Output plain text only, no JSON.
    """)
    raw = call_ollama(model, prompt)
    return raw

def write_docx(resume_text: str, output_path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for para in resume_text.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(output_path)

# ------------------------------------------------------------------
# 5. Main driver
# ------------------------------------------------------------------
def main(job_ad_path: Path, resume_path: Path, models: list[str]):
    job_ad_raw = read_text_file(job_ad_path)
    if resume_path.suffix.lower() == ".pdf":
        resume_raw = read_pdf(resume_path)
    else:
        resume_raw = read_docx(resume_path)

    job_ad_text = clean_text(job_ad_raw)
    resume_text = clean_text(resume_raw)

    for model in models:
        print(f"\n=== Using Ollama model: {model} ===")

        # Summarise
        job_summary = summarise_job_ad(job_ad_text, model)
        print("Job title:", job_summary["title"])
        print(f"Extracted {len(job_summary['key_skills'])} key skills")

        # Extract résumé skills
        resume_skills = extract_resume_skills(resume_text, model)
        print(f"Found {len(resume_skills)} distinct skills in résumé")

        # Align
        alignment = align_skills(job_summary["key_skills"], resume_skills)

        # Generate customised résumé
        customised = generate_custom_resume(job_summary, resume_text, alignment, model)

        # Output
        out_file = resume_path.stem + f"_updated_{model.replace('-', '_')}.docx"
        write_docx(customised, Path(out_file))
        print(f"Saved to {out_file}")

# ------------------------------------------------------------------
# 6. CLI
# ------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Customise a résumé using Ollama models")
    parser.add_argument("--job_ad", required=True, type=Path,
                        help="Plain‑text file containing the job posting")
    parser.add_argument("--resume", required=True, type=Path,
                        help="Applicant résumé (PDF or DOCX)")
    parser.add_argument("--models", required=True, type=str,
                        help="Comma‑separated list of Ollama model names (e.g. llama2, phi3)")
    args = parser.parse_args()

    model_list = [m.strip() for m in args.models.split(",") if m.strip()]
    main(args.job_ad, args.resume, model_list)